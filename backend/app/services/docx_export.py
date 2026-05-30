"""Render a `TailoredResume` to a DOCX byte stream, ATS-standard formatting.

Built from the shared `resume_layout` block list, so the DOCX and the PDF
(`pdf_export`) carry IDENTICAL text — only the styling differs.

ATS conventions:
  - Single column. No tables, columns, text boxes, images, or headers/footers.
  - ATS-safe font (Calibri). Body ~10.5pt, name ~20pt. 0.6" margins.
  - Closed-list section headings (Professional Summary / Skills / Experience /
    Education / Projects / Certifications).
  - Two-line entry blocks: bold line 1 (title/degree), light line 2
    (company/institution + dates).
  - Visual mode: a horizontal rule under each heading; dates flush-right via
    a right-aligned TAB STOP (a real `<w:tab w:val="right">`, NOT a table).
  - Plain mode: no rules, headings as plain bold text, dates inline after the
    company/institution separated by " | ".

The 2-page hard cap is enforced upstream in `tailor.generate_resume` (it
measures the rendered PDF and trims), so this renderer just lays out the
content it's given.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.services.resume_layout import (
    SEP,
    Bullet,
    Entry,
    Header,
    Heading,
    Para,
    build_blocks,
)
from app.services.tailor import TailoredResume

# Type styling constants — kept here so the whole document reads from one place.
_FONT = "Calibri"
_BODY_PT = 10.5
_NAME_PT = 20.0
_HEADLINE_PT = 12.0
_SMALL_PT = 9.5
_LIGHT = RGBColor(0x55, 0x55, 0x55)
_HEADING_INK = RGBColor(0x22, 0x22, 0x22)
_MARGIN_INCHES = 0.6

# Header alignment is a user choice (left/center/right), applied to the name +
# headline + contact/links block ONLY. Body text (summary and everything below)
# stays left-aligned regardless — readable body copy is always left.
_DOCX_HEADER_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def render_docx(
    resume: TailoredResume,
    candidate: dict[str, Any] | None = None,  # noqa: ARG001 — kept for call-site compatibility
    *,
    mode: str | None = None,
    header_alignment: str = "center",
) -> bytes:
    """Render `resume` to DOCX bytes. `mode` is "visual" (default) or
    "plain"; falls back to the resume's own `meta.mode`. `header_alignment`
    ("left" | "center" | "right", default "center") positions the name +
    contact header block; body text always stays left. The legacy
    `candidate` arg is accepted and ignored — contact now lives on the
    resume itself (reconciled server-side from the profile)."""
    chosen = (mode or resume.meta.mode or "visual").lower()
    plain = chosen == "plain"
    align = _DOCX_HEADER_ALIGN.get(header_alignment, WD_ALIGN_PARAGRAPH.CENTER)
    blocks = build_blocks(resume)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = _inches(_MARGIN_INCHES)
    section.right_margin = _inches(_MARGIN_INCHES)
    section.top_margin = _inches(_MARGIN_INCHES)
    section.bottom_margin = _inches(_MARGIN_INCHES)
    right_tab_pos = section.page_width - section.left_margin - section.right_margin

    style = doc.styles["Normal"]
    style.font.name = _FONT
    style.font.size = Pt(_BODY_PT)
    style.paragraph_format.space_after = Pt(2)

    for block in blocks:
        if isinstance(block, Header):
            _render_header(doc, block, alignment=align)
        elif isinstance(block, Heading):
            _render_heading(doc, block.text, plain=plain)
        elif isinstance(block, Para):
            doc.add_paragraph(block.text)
        elif isinstance(block, Entry):
            _render_entry(doc, block, right_tab_pos=right_tab_pos, plain=plain)
        elif isinstance(block, Bullet):
            _render_bullet(doc, block.text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- block renderers --------------------------------------------------------


def _render_header(doc: Any, header: Header, *, alignment: Any = WD_ALIGN_PARAGRAPH.CENTER) -> None:
    name_p = doc.add_paragraph()
    name_p.alignment = alignment
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(header.name)
    name_run.bold = True
    name_run.font.size = Pt(_NAME_PT)

    if header.headline:
        hp = doc.add_paragraph()
        hp.alignment = alignment
        hr = hp.add_run(header.headline)
        hr.font.size = Pt(_HEADLINE_PT)
        hr.font.color.rgb = _LIGHT

    for line in (header.contact_line, header.links_line):
        if line:
            cp = doc.add_paragraph(line)
            cp.alignment = alignment
            cp.runs[0].font.size = Pt(_SMALL_PT)
            cp.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _render_heading(doc: Any, text: str, *, plain: bool) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    if not plain:
        # Visual mode: a touch of ink + a hairline rule under the heading.
        run.font.color.rgb = _HEADING_INK
        _add_bottom_border(p)


def _render_entry(doc: Any, entry: Entry, *, right_tab_pos: int, plain: bool) -> None:
    # Line 1 — bold (title / degree / project / cert name).
    line1_p = doc.add_paragraph()
    line1_p.paragraph_format.space_after = Pt(0)
    if entry.line1:
        line1_p.add_run(entry.line1).bold = True

    # Line 2 — light (company/institution + dates). Skip entirely when both
    # halves are empty (e.g. a project with only a name).
    if not entry.left and not entry.right:
        return

    line2_p = doc.add_paragraph()
    line2_p.paragraph_format.space_after = Pt(4)

    if plain:
        # Dates inline after the left text, separated by " | ".
        text = SEP.join(part for part in (entry.left, entry.right) if part)
        run = line2_p.add_run(text)
        run.font.size = Pt(_SMALL_PT)
        run.font.color.rgb = _LIGHT
    else:
        # Visual: left text, then dates flush-right via a right tab stop.
        if right_tab_pos:
            line2_p.paragraph_format.tab_stops.add_tab_stop(right_tab_pos, WD_TAB_ALIGNMENT.RIGHT)
        if entry.left:
            lr = line2_p.add_run(entry.left)
            lr.font.size = Pt(_SMALL_PT)
            lr.font.color.rgb = _LIGHT
        if entry.right:
            line2_p.add_run("\t")
            rr = line2_p.add_run(entry.right)
            rr.font.size = Pt(_SMALL_PT)
            rr.italic = True
            rr.font.color.rgb = _LIGHT


def _render_bullet(doc: Any, text: str) -> None:
    """A bullet rendered as a normal paragraph with a leading "- " and a
    hanging indent. We use a plain hyphen (not the List Bullet glyph) so the
    text is ATS-clean and byte-identical to the PDF bullet."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"- {text}")


# --- low-level helpers ------------------------------------------------------


def _inches(value: float) -> int:
    from docx.shared import Inches  # noqa: PLC0415

    return Inches(value)


def _add_bottom_border(paragraph: Any) -> None:
    """Add a thin bottom border to a paragraph — the section-heading rule."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


__all__ = ["render_docx"]
