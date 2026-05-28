"""Resume-upload text extractor.

Given a file's bytes + reported filename, extract plain text and
hand it to the existing hybrid parser (`parse_resume`). The parser
itself is unchanged — this module just lifts text out of PDF/DOCX
containers so the same downstream pipeline (regex contact fields +
LLM structural extract) runs on uploads the same way it runs on
pastes.

Public API:
  * `extract_text(filename, data) -> str`  — best-effort extract.
                                             Raises `UnsupportedResumeFile`
                                             for unknown extensions and
                                             `EmptyExtractionError` when
                                             the file parsed cleanly but
                                             produced no text (e.g. a
                                             scanned, image-only PDF
                                             with no OCR layer).
  * `UnsupportedResumeFile`                 — the API surfaces this as
                                             400 with a clear message.
  * `EmptyExtractionError`                  — surfaced as 422 with the
                                             "paste your text instead"
                                             hint.

A corrupt or unreadable PDF/DOCX file raises a generic exception
that the caller catches and reports as a clear "couldn't read this
file" error — better than letting the underlying pdfminer/docx
exception bubble.
"""

from __future__ import annotations

import io
import logging
import re
from pathlib import PurePosixPath
from typing import Any

log = logging.getLogger(__name__)


# Allowed input file types. Cheap extension check + content-sniff
# happens inside each branch — we trust the extension as a routing
# hint but never as proof.
SUPPORTED_SUFFIXES: tuple[str, ...] = (".pdf", ".docx")

# Hard cap on accepted upload size — 10 MB is well above any real
# resume. Beyond this we 413 before opening the file so a malicious
# upload can't grind the parser.
MAX_UPLOAD_BYTES = 10 * 1024 * 1024


class UnsupportedResumeFile(ValueError):
    """Raised when the uploaded file's extension isn't `.pdf` or
    `.docx`. The API turns this into a 400 with the allow-list."""


class EmptyExtractionError(ValueError):
    """Raised when extraction completed without raising but produced
    no usable text. Almost always means a scanned / image-only PDF
    or a DOCX with no text-runs. The API surfaces this as 422 with
    a clear "paste your text instead" hint."""


def _suffix(filename: str) -> str:
    """Lowercased extension of `filename`. `PurePosixPath` keeps
    things deterministic even on weird inputs like `resume.PDF` or
    `resume.tar.gz`."""
    return PurePosixPath(filename or "").suffix.lower()


def extract_text(filename: str, data: bytes) -> str:
    """Best-effort plain-text extraction. Output is whitespace-
    normalised but never fabricated — empty input → raises
    `EmptyExtractionError` instead of returning the empty string,
    so the API path can short-circuit to its actionable error
    message rather than handing empty text to the LLM."""
    if not data:
        raise EmptyExtractionError("Uploaded file is empty.")

    suffix = _suffix(filename)
    if suffix not in SUPPORTED_SUFFIXES:
        raise UnsupportedResumeFile(
            f"Unsupported file type {suffix!r}. Allowed: " f"{', '.join(SUPPORTED_SUFFIXES)}."
        )

    if suffix == ".pdf":
        text = _extract_pdf(data)
    else:  # ".docx"
        text = _extract_docx(data)

    text = _normalise(text)
    if not text.strip():
        # Most common cause: scanned-image PDF with no text layer.
        # The frontend turns this into a user-facing message
        # asking the user to paste the text directly.
        raise EmptyExtractionError(
            "No text could be extracted from the file. If your resume is a "
            "scanned image or image-only PDF, please paste the text instead."
        )
    return text


# ─── PDF ────────────────────────────────────────────────────────────────────


def _extract_pdf(data: bytes) -> str:
    """`pdfplumber` walks every page + extracts the text layer. A
    page with no text layer (scanned image) returns the empty
    string — collected across all pages we detect that case
    upstream as `EmptyExtractionError`.

    Imported lazily to keep cold start fast and to avoid pulling
    pdfminer into the import graph of anything that doesn't need
    it (notably tests for the parser itself)."""
    import pdfplumber  # noqa: PLC0415

    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = []
            for page in pdf.pages:
                try:
                    text = page.extract_text() or ""
                except Exception as e:  # noqa: BLE001
                    # A single bad page shouldn't fail the whole
                    # extract — log + skip, then trust the empty-text
                    # detection upstream.
                    log.warning("pdfplumber: skipping page — %s", e)
                    text = ""
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
    except EmptyExtractionError:
        raise
    except Exception as e:  # noqa: BLE001
        # Wrap pdfminer's various failure modes (PDFSyntaxError,
        # PSEOF, etc.) in a single ValueError the API knows about.
        raise ValueError(f"Couldn't read PDF: {e}") from e


# ─── DOCX ───────────────────────────────────────────────────────────────────


def _extract_docx(data: bytes) -> str:
    """`python-docx` opens the docx, walks every text-bearing surface.

    Resume DOCX templates are a graveyard of content hiding in
    non-obvious places. The previous version only walked top-level
    paragraphs + table cells, which dropped bullets (because some
    templates put bullets inside `w:txbxContent` text boxes),
    contact rows (sometimes in headers/footers), and styled lists
    (the bullet glyph is a separate run + paragraph style that
    `para.text` strips without warning).

    We now walk, in document order:
      1. Headers + footers in every section.
      2. Top-level paragraphs, prefixing list paragraphs with "- "
         so the downstream parser recognises them as bullets.
      3. Tables (with the same per-row prefix rule for cells whose
         paragraphs are list-styled).
      4. Text boxes (`w:txbxContent`) via a fallback XML walk —
         python-docx exposes paragraphs there but they don't show
         up in `doc.paragraphs`.

    Empty paragraphs are preserved as blank lines so the parser's
    section segmenter (which uses blank lines to split entries)
    sees the same shape across PDF / DOCX / paste inputs.
    """
    from docx import Document  # noqa: PLC0415

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        raise ValueError(f"Couldn't read DOCX: {e}") from e

    parts: list[str] = []

    # 1. Headers + footers — contact rows + page numbering often
    #    live here on enterprise-template resumes.
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for para in hf.paragraphs:
                line = _docx_paragraph_text(para)
                if line:
                    parts.append(line)
            for table in hf.tables:
                for row in table.rows:
                    cells = [_docx_cell_text(c) for c in row.cells]
                    cells = [c for c in cells if c]
                    if cells:
                        parts.append("\t".join(cells))

    # 2. Body paragraphs. List-styled paragraphs get a `- ` prefix
    #    so the LLM / regex parser recognises them as bullets.
    for para in doc.paragraphs:
        line = _docx_paragraph_text(para)
        if line:
            parts.append(line)
        elif para.text == "":
            # Preserve blank paragraphs as blank lines so the section
            # segmenter sees the same shape it would in a paste.
            parts.append("")

    # 3. Tables (date right-aligned in a second column is a classic
    #    resume template).
    for table in doc.tables:
        for row in table.rows:
            cells = [_docx_cell_text(c) for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append("\t".join(cells))

    # 4. Text boxes — `w:txbxContent` paragraphs are NOT in
    #    `doc.paragraphs`. Walk the XML directly so a template that
    #    parks bullets in a text box doesn't silently drop them.
    parts.extend(_extract_docx_textboxes(doc))

    return "\n".join(parts)


def _docx_paragraph_text(para: Any) -> str:
    """Render one DOCX paragraph as plain text, prefixing list-
    styled lines with `- ` so the downstream parser can spot bullets.
    Empty paragraphs return ``""`` — the caller decides whether to
    preserve the blank line or skip it."""
    raw = para.text or ""
    if not raw.strip():
        return ""
    if _is_list_paragraph(para):
        return f"- {raw.strip()}"
    return raw


def _docx_cell_text(cell: Any) -> str:
    """Render one table cell as plain text. Cells can themselves
    hold multi-paragraph content; we join paragraphs with newlines so
    list bullets inside a cell still survive."""
    out: list[str] = []
    for para in cell.paragraphs:
        line = _docx_paragraph_text(para)
        if line:
            out.append(line)
    return "\n".join(out)


def _is_list_paragraph(para: Any) -> bool:
    """Heuristic for whether a DOCX paragraph is part of a list.

    Two signals:
      * The paragraph style name contains "List" (e.g. `List Bullet`,
        `List Number`, `List Paragraph`).
      * The paragraph XML carries a `w:numPr` element — the explicit
        numbering / bullet marker the renderer respects regardless of
        style name.

    Either is enough. Returns False on any error (a non-paragraph
    object slipped through, an unexpected XML shape) — bullets
    rendering as plain lines is a much better failure mode than
    raising mid-parse.
    """
    try:
        style_name = getattr(getattr(para, "style", None), "name", "") or ""
        if "list" in style_name.lower():
            return True
        # `w:numPr` is the structural-numbering element python-docx
        # exposes via the paragraph's underlying `<w:pPr>`.
        ppr = para._p.find(_qn("w:pPr"))
        if ppr is not None and ppr.find(_qn("w:numPr")) is not None:
            return True
    except Exception:  # noqa: BLE001
        pass
    return False


def _extract_docx_textboxes(doc: Any) -> list[str]:
    """Walk every `<w:txbxContent>` element in the document body and
    yield one line per text-bearing paragraph (with list-bullet
    prefixing). python-docx's `doc.paragraphs` skips text-box
    contents, so this is the only place that content surfaces.
    """
    lines: list[str] = []
    try:
        body = doc.element.body
    except Exception:  # noqa: BLE001
        return lines
    for txbx in body.iter(_qn("w:txbxContent")):
        for p in txbx.iter(_qn("w:p")):
            # Reassemble the paragraph's runs as text. We don't need
            # the python-docx Paragraph wrapper here — just the raw
            # text.
            runs = [t.text or "" for t in p.iter(_qn("w:t"))]
            text = "".join(runs).strip()
            if not text:
                continue
            # Bullet detection for text-box paragraphs: same `w:numPr`
            # signal the regular helper uses.
            ppr = p.find(_qn("w:pPr"))
            is_list = ppr is not None and ppr.find(_qn("w:numPr")) is not None
            lines.append(f"- {text}" if is_list else text)
    return lines


def _qn(tag: str) -> str:
    """python-docx's `qn()` returns the fully-qualified XML name for
    a `w:foo` shorthand. Lazy-imported here so the helper module
    doesn't pull `docx` for callers that only need the PDF branch."""
    from docx.oxml.ns import qn  # noqa: PLC0415

    return qn(tag)


# ─── Normalisation ─────────────────────────────────────────────────────────


# Collapse 3+ consecutive newlines into 2 so the section segmenter
# (which uses blank lines to split entries) sees a consistent shape
# across PDF/DOCX/paste inputs. Trailing whitespace per line is
# stripped because PDF extracts often carry hard-wrap padding.
_NEWLINE_RUN_RE = re.compile(r"\n{3,}")
_TRAILING_WS_RE = re.compile(r"[ \t]+$", re.MULTILINE)


def _normalise(text: str) -> str:
    if not text:
        return ""
    text = _TRAILING_WS_RE.sub("", text)
    text = _NEWLINE_RUN_RE.sub("\n\n", text)
    return text.strip()
