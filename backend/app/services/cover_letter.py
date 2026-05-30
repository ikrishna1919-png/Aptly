"""Cover-letter generation + rendering for the /ats hub.

Reuses the tailor Anthropic client + JSON-extraction + sanitizer, with a
cover-letter-specific prompt and a much simpler output shape (business-letter
prose, NOT the resume block model). Same hard anti-fabrication rule as resume
tailoring: every claim is grounded in the user's real profile.
"""

from __future__ import annotations

import io
import json
from typing import Any

from pydantic import BaseModel, Field

from app.config import Settings, get_settings
from app.services.demo_candidate import get_candidate
from app.services.tailor import (
    MODEL,
    _build_client,
    _extract_json_object,
    _sanitize_obj,
)

# 3 pre-built cover-letter formats (fewer than resume — letters vary less).
COVER_FORMATS = ("traditional", "modern", "concise")


class CoverLetterContent(BaseModel):
    """Structured letter so DOCX + PDF share one source of truth."""

    date: str = ""
    recipient: str = ""  # e.g. "Hiring Team, Acme Corp"
    greeting: str = "Dear Hiring Manager,"
    paragraphs: list[str] = Field(default_factory=list)
    closing: str = "Sincerely,"
    signature: str = ""


_TONE = {
    "formal": "formal and polished",
    "confident": "confident and results-forward",
    "warm": "warm and personable (still professional)",
}
_LENGTH = {"short": "about 200 words", "standard": "about 350 words", "long": "about 500 words"}
_OPENING = {
    "value": "open by leading with the candidate's most relevant value to this role",
    "story": "open with a brief, true story or moment from the candidate's real experience",
    "question": "open with an engaging question that frames the candidate's fit",
}

_SYSTEM = (
    "You write job-application cover letters. You are given the CANDIDATE "
    "PROFILE (real experience), the TARGET JOB, the company, and style "
    "preferences. Write a tailored letter.\n\n"
    "NON-NEGOTIABLE: NEVER invent achievements, metrics, certifications, "
    "employers, or roles. Ground every claim in the candidate profile. Mirror "
    "relevant keywords from the job description naturally — no stuffing. No "
    "en/em dashes; use a hyphen or rewrite. Straight quotes only.\n\n"
    'Return ONLY JSON: {"date": "", "recipient": "", "greeting": "", '
    '"paragraphs": ["", ""], "closing": "Sincerely,", "signature": ""}. '
    "No markdown fences, no prose outside the JSON. The first char is { and "
    "the last is }."
)


def _instructions(questions: dict[str, Any] | None, company: str, hook: str) -> str:
    q = questions or {}
    tone = _TONE.get(str(q.get("tone") or "confident").lower(), _TONE["confident"])
    length = _LENGTH.get(str(q.get("length") or "standard").lower(), _LENGTH["standard"])
    opening = _OPENING.get(str(q.get("opening") or "value").lower(), _OPENING["value"])
    extra = str(q.get("additional") or "").strip()
    lines = [
        f"- Tone: {tone}.",
        f"- Length: {length}.",
        f"- Opening: {opening}.",
    ]
    if company:
        lines.append(f"- Company: {company}.")
    if hook:
        lines.append(f"- Incorporate this true detail the candidate provided: {hook}")
    if extra:
        lines.append(f"- Additional instruction from the candidate: {extra}")
    return "\n".join(lines)


def _demo_letter(candidate: dict[str, Any], company: str) -> CoverLetterContent:
    name = candidate.get("name") or "Your Name"
    summary = candidate.get("summary") or "an experienced professional"
    co = company or "your team"
    return CoverLetterContent(
        recipient=f"Hiring Team{f', {company}' if company else ''}",
        greeting="Dear Hiring Manager,",
        paragraphs=[
            f"I am writing to express my interest in joining {co}. As {summary}, "
            "I believe my background aligns well with what you are looking for.",
            "[demo mode] This is placeholder text. With an Anthropic API key "
            "configured, this letter is generated from your real profile and the "
            "job description, grounded only in your actual experience.",
            "Thank you for your consideration. I would welcome the chance to discuss "
            "how I can contribute.",
        ],
        closing="Sincerely,",
        signature=name,
    )


def generate_cover_letter(
    db: Any,
    *,
    user_id: int | None,
    jd_text: str,
    company_name: str,
    hook: str,
    questions: dict[str, Any] | None,
    settings: Settings | None = None,
    client: Any | None = None,
) -> CoverLetterContent:
    """Generate a grounded cover letter. Demo fallback when no key."""
    settings = settings or get_settings()
    candidate = get_candidate(db, user_id=user_id)

    if not settings.has_anthropic_key:
        return _demo_letter(candidate, company_name)

    api = _build_client(settings, client)
    user_content = (
        "CANDIDATE PROFILE (do not modify these facts):\n"
        + json.dumps(candidate, indent=2, sort_keys=True)
        + f"\n\nTARGET JOB:\n{jd_text.strip()}\n\n"
        + _instructions(questions, company_name, hook)
        + "\n\nReturn ONLY the cover letter JSON."
    )
    resp = api.messages.create(
        model=MODEL,
        max_tokens=1500,
        system=[{"type": "text", "text": _SYSTEM}],
        messages=[{"role": "user", "content": user_content}],
    )
    text = next(b.text for b in resp.content if getattr(b, "type", None) == "text")
    obj = _extract_json_object(text)
    counter: dict[str, int] = {}
    obj = _sanitize_obj(obj, counter)  # strip en/em dashes + smart quotes
    letter = CoverLetterContent.model_validate(obj)
    # Reconcile signature from the authoritative profile.
    if not letter.signature:
        letter.signature = candidate.get("name") or ""
    return letter


# ─── Rendering (simple business letter — no resume block model) ──────────────


def _letter_lines(letter: CoverLetterContent) -> list[tuple[str, bool]]:
    """(text, is_blank) sequence shared by both renderers."""
    out: list[tuple[str, bool]] = []
    if letter.date:
        out.append((letter.date, False))
        out.append(("", True))
    if letter.recipient:
        out.append((letter.recipient, False))
        out.append(("", True))
    out.append((letter.greeting or "Dear Hiring Manager,", False))
    out.append(("", True))
    for p in letter.paragraphs:
        out.append((p, False))
        out.append(("", True))
    out.append((letter.closing or "Sincerely,", False))
    out.append((letter.signature or "", False))
    return out


def render_cover_docx(letter: CoverLetterContent, fmt: str = "traditional") -> bytes:
    from docx import Document  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    serif = fmt in ("traditional",)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Georgia" if serif else "Calibri"
    style.font.size = Pt(11)
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = _docx_inch(1.0)
        sec.top_margin = sec.bottom_margin = _docx_inch(1.0)
    for text, blank in _letter_lines(letter):
        if blank:
            doc.add_paragraph("")
        else:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_inch(v: float):  # noqa: ANN202
    from docx.shared import Inches  # noqa: PLC0415

    return Inches(v)


def render_cover_pdf(letter: CoverLetterContent, fmt: str = "traditional") -> bytes:
    from xml.sax.saxutils import escape  # noqa: PLC0415

    from reportlab.lib.enums import TA_LEFT  # noqa: PLC0415
    from reportlab.lib.pagesizes import letter as letter_size  # noqa: PLC0415
    from reportlab.lib.styles import ParagraphStyle  # noqa: PLC0415
    from reportlab.lib.units import inch  # noqa: PLC0415
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # noqa: PLC0415

    serif = fmt in ("traditional",)
    body = ParagraphStyle(
        "body",
        fontName="Times-Roman" if serif else "Helvetica",
        fontSize=11,
        leading=15,
        alignment=TA_LEFT,
    )
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter_size,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title="Cover letter",
    )
    flow: list = []
    for text, blank in _letter_lines(letter):
        if blank:
            flow.append(Spacer(1, 8))
        else:
            flow.append(Paragraph(escape(text), body))
    doc.build(flow)
    return buf.getvalue()
