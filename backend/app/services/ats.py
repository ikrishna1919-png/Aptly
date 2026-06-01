"""ATS resume-hub generation logic.

Sits on top of the PR #58/#66/#69 tailor machinery (prompt-based JSON, the
`GeneratedResume` schema, the sanitizer, contact reconciliation, the page
measurer) and adds the /ats-specific bits:

  * `build_customization_addendum` — turns the 6 customization answers into
    prompt instructions.
  * `generate_ats` — generate a tailored `TailoredResume` from a pasted JD
    (no Job row) + customization, streaming, capped to the chosen page count.
  * `compute_keyword_edits` / `apply_docx_edits` — the Option-2 DOCX
    keyword-injection path: ask the model for minimal {original→replacement}
    swaps and apply them to the uploaded DOCX run-by-run, preserving format.

Kept separate from `tailor.py` so the heavily-tested tailor functions are
untouched; we import their building blocks rather than fork them.
"""

from __future__ import annotations

import io
import json
import logging
import time
from collections.abc import Callable
from typing import TYPE_CHECKING, Any

from app.config import Settings, get_settings
from app.services.demo_candidate import get_candidate
from app.services.tailor import (
    _STREAM_SNAPSHOT_INTERVAL_SECONDS,
    _SYSTEM_GENERATE_JSON,
    MODEL,
    GeneratedResume,
    ResumeMeta,
    TailoredResume,
    _build_client,
    _demo_resume,
    _extract_json_object,
    _measure_pages,
    _reconcile_contact,
    _system_blocks,
    _truncate_to_two_pages,
    loads_partial,
    sanitize_generated,
)

if TYPE_CHECKING:
    from sqlalchemy.orm import Session

log = logging.getLogger(__name__)


# ─── Customization → prompt ─────────────────────────────────────────────────

_TONE = {
    "formal": "Use a formal, polished register.",
    "confident": "Use a confident, results-forward register.",
    "conversational": "Use a clear, natural, conversational register (still professional).",
}
_EMPHASIS = {
    "technical": "Foreground technical depth — tools, systems, and engineering detail.",
    "leadership": "Foreground leadership impact — ownership, mentoring, cross-team influence.",
    "execution": "Foreground project execution — delivery, scope, and outcomes shipped.",
    "mixed": "Keep a balanced emphasis across technical depth, leadership, and execution.",
}


def build_customization_addendum(answers: dict[str, Any] | None) -> str:
    """Render the 6 customization answers into an instruction block appended to
    the generate prompt. Empty/absent answers fall back to sensible defaults
    (and skills/roles default to 'AI chooses')."""
    a = answers or {}
    lines: list[str] = ["\n\nCUSTOMIZATION (honor these without inventing anything):"]

    length = str(a.get("length") or "1").strip()
    pages = 2 if length in ("2", "2 pages", "two") else 1
    lines.append(f"- Target length: {pages} page{'s' if pages == 2 else ''}. Be concise.")

    tone = _TONE.get(str(a.get("tone") or "confident").lower())
    if tone:
        lines.append(f"- Tone: {tone}")

    emphasis = _EMPHASIS.get(str(a.get("emphasis") or "mixed").lower())
    if emphasis:
        lines.append(f"- Emphasis: {emphasis}")

    skills = [s for s in (a.get("skills") or []) if str(s).strip()]
    if skills:
        lines.append(
            "- Foreground these skills where the candidate genuinely has them: "
            + ", ".join(skills)
            + ". Do NOT add skills the profile lacks."
        )

    roles = [r for r in (a.get("roles") or []) if str(r).strip()]
    if roles:
        lines.append("- Give more weight to these roles: " + ", ".join(roles) + ".")

    extra = str(a.get("additional") or "").strip()
    if extra:
        lines.append(f"- Additional instructions from the candidate: {extra}")

    return "\n".join(lines)


def _target_pages(answers: dict[str, Any] | None) -> int:
    length = str((answers or {}).get("length") or "1").strip()
    return 2 if length in ("2", "2 pages", "two") else 1


def _jd_block(jd_text: str, *, job_title: str | None = None) -> str:
    title = (job_title or "").strip() or "(see description)"
    return f"TARGET JOB:\nTitle: {title}\n\nJob description:\n{jd_text.strip()}"


def _cap_to_pages(resume: TailoredResume, target: int) -> TailoredResume:
    """Reuse the deterministic 2-page trimmer; for a 1-page target we still
    only hard-trim past 2 pages (never butcher content to force 1 page — the
    prompt asks for 1, this is just the safety ceiling)."""
    pages = _measure_pages(resume)
    if pages > 2:
        resume = _truncate_to_two_pages(resume)
        pages = _measure_pages(resume)
    resume.meta.pages_estimate = max(1, min(2, pages))
    return resume


# ─── Generation (JD-paste / PDF-fallback paths) ─────────────────────────────


def generate_ats(
    db: Session,
    *,
    user_id: int | None,
    jd_text: str,
    customization: dict[str, Any] | None,
    settings: Settings | None = None,
    client: Any | None = None,
    stream_cb: Callable[[GeneratedResume], None] | None = None,
    deadline: float | None = None,
) -> TailoredResume:
    """Generate a tailored resume from a pasted JD + customization answers.

    Mirrors `tailor.generate_resume` but is JD-text driven (no Job row) and
    threads the 6 customization answers into the prompt. Streams partial
    snapshots via `stream_cb` when a real key is configured; falls back to the
    deterministic demo resume otherwise."""
    settings = settings or get_settings()
    candidate = get_candidate(db, user_id=user_id)
    target = _target_pages(customization)

    if not settings.has_anthropic_key:
        stub = type("J", (), {"skills": [], "title": "", "company": "", "description": jd_text})()
        gen = _demo_resume(stub, {}, candidate=candidate)
        gen = sanitize_generated(gen)
        gen = _reconcile_contact(gen, candidate)
        return _cap_to_pages(
            TailoredResume(**gen.model_dump(), meta=ResumeMeta(mode="visual")), target
        )

    user_content = (
        _jd_block(jd_text)
        + build_customization_addendum(customization)
        + "\n\nReturn ONLY the tailored resume as a single JSON object in the structure "
        "given in the system prompt. Never invent facts not present in the candidate profile."
    )
    gen = _stream_generated(
        system=_system_blocks(_SYSTEM_GENERATE_JSON, candidate),
        user_content=user_content,
        client=client,
        settings=settings,
        on_partial=stream_cb,
        deadline=deadline,
    )
    gen = sanitize_generated(gen)
    gen = _reconcile_contact(gen, candidate)
    return _cap_to_pages(TailoredResume(**gen.model_dump(), meta=ResumeMeta(mode="visual")), target)


def _stream_generated(
    *,
    system: list[dict[str, Any]],
    user_content: str,
    client: Any | None,
    settings: Settings,
    on_partial: Callable[[GeneratedResume], None] | None,
    deadline: float | None,
) -> GeneratedResume:
    """Stream a `GeneratedResume` from the Messages API (prompt-based JSON, no
    grammar). Emits throttled partial snapshots; one correction retry on a
    final JSON syntax error. Mirrors tailor's streaming contract."""
    api = _build_client(settings, client)
    kwargs: dict[str, Any] = {
        "model": MODEL,
        "max_tokens": 4000,
        "system": system,
        "messages": [{"role": "user", "content": user_content}],
    }
    chunks: list[str] = []
    last_emit = 0.0
    with api.messages.stream(**kwargs) as stream:
        for delta in stream.text_stream:
            chunks.append(delta)
            now = time.monotonic()
            if deadline is not None and now > deadline:
                raise TimeoutError("ats generation exceeded the time budget")
            if on_partial is not None and (now - last_emit) >= _STREAM_SNAPSHOT_INTERVAL_SECONDS:
                last_emit = now
                partial = loads_partial("".join(chunks))
                if partial:
                    try:
                        on_partial(GeneratedResume.model_validate(partial))
                    except Exception:  # noqa: BLE001
                        pass
        final = stream.get_final_message()
        text = "".join(chunks) or _first_text(final)
    try:
        obj = _extract_json_object(text)
    except (json.JSONDecodeError, ValueError):
        retry = api.messages.create(
            model=MODEL,
            max_tokens=4000,
            system=system,
            messages=[
                {"role": "user", "content": user_content},
                {"role": "assistant", "content": text},
                {
                    "role": "user",
                    "content": "Your previous reply was not valid JSON. Reply again with "
                    "ONLY the corrected JSON object — no fences, no prose.",
                },
            ],
        )
        obj = _extract_json_object(_first_text(retry))
    return GeneratedResume.model_validate(obj)


def _first_text(response: Any) -> str:
    for block in response.content:
        if getattr(block, "type", None) == "text":
            return block.text
    raise RuntimeError("Anthropic response contained no text block")


# ─── Option 2: DOCX keyword injection ───────────────────────────────────────


class DocxEdit(dict):
    pass


_KEYWORD_SYSTEM = (
    "You optimize an existing resume for a target job by proposing the MINIMUM "
    "set of in-place wording swaps that weave in the job's keywords TRUTHFULLY. "
    "You are NOT rewriting the resume. Each edit replaces an exact existing "
    "phrase with a lightly-adjusted version that incorporates a relevant JD "
    "keyword the candidate genuinely supports. NEVER invent skills, employers, "
    "metrics, or experience. Each replacement_text MUST stay CLOSE to the "
    "original's length (within a few characters) — do NOT materially grow a "
    "bullet or summary line, so the resume's line count and layout stay visually "
    "intact. Swap a word or two; never expand a phrase into a sentence.\n\n"
    'Return ONLY JSON: {"edits": [{"original_text": "...", '
    '"replacement_text": "...", "reason": "..."}]}. `original_text` MUST '
    "be copied verbatim from the resume so it can be found exactly. Propose at "
    "most 12 edits.\n\n"
    "Return ONLY valid JSON. No markdown formatting. No commentary before or "
    "after the JSON. The response must be parseable with Python's json.loads() "
    "with no preprocessing."
)

# Corrective turn appended on the one retry when the first reply won't parse.
_KEYWORD_JSON_FIX = (
    "Your previous response had a JSON syntax error: {error}. Return ONLY the "
    "corrected JSON, no markdown fences, no prose."
)


class KeywordInjectionError(Exception):
    """Raised when keyword-edit JSON can't be parsed after a retry. The worker
    turns this into a clear, user-facing message (never a raw traceback)."""


# Monitoring: how often the keyword-edit JSON failed to parse, by stage. Read
# via `keyword_parse_failures` so recurrence can be watched after the fix.
keyword_parse_failures: dict[str, int] = {"first": 0, "retry": 0}


def _keyword_addendum(answers: dict[str, Any] | None) -> str:
    """Fold the option-B gap/relevance answers into a STEER block for the
    keyword-edit prompt. They steer WHICH existing phrases to rewrite and how
    (relevance, JD skills the candidate confirmed, credible metrics) — they
    NEVER add content. Empty answers → empty string (no steer)."""
    a = answers or {}
    lines: list[str] = []

    # Yes/No gap confirmations: only "yes" entries steer rewrites; details (if
    # given) sharpen WHICH existing phrase to reword. "no" is ignored.
    confirmed: list[str] = []
    for g in a.get("gaps") or []:
        if isinstance(g, dict) and str(g.get("answer", "")).strip().lower() == "yes":
            q = str(g.get("question") or g.get("skill") or "").strip()
            d = str(g.get("details") or "").strip()
            if q:
                confirmed.append(f"{q} ({d})" if d else q)
    if confirmed:
        lines.append(
            "- Skills/experience the candidate CONFIRMED they genuinely have — reword "
            "EXISTING phrases to surface these where truthful (never add new lines): "
            + "; ".join(confirmed)
            + "."
        )

    missing_exp = str(a.get("missing_experience") or "").strip()
    if missing_exp:
        lines.append(
            "- Relevant experience the candidate has done but that isn't clearly on the "
            "resume — use ONLY to decide which EXISTING phrases to strengthen/reword toward "
            f"the JD (do NOT add it as new content): {missing_exp}"
        )
    skills = [s for s in (a.get("skills") or []) if str(s).strip()]
    if skills:
        lines.append(
            "- JD skills the candidate confirmed they genuinely have — prefer rewording "
            "existing phrases to surface these where truthful: " + ", ".join(skills) + "."
        )
    roles = [r for r in (a.get("roles") or []) if str(r).strip()]
    if roles:
        lines.append(
            "- Foreground wording in the candidate's most-relevant roles: " + ", ".join(roles) + "."
        )
    metrics = str(a.get("metrics") or "").strip()
    if metrics:
        lines.append(
            "- Credible numbers for the candidate's most relevant work — weave into EXISTING "
            f"phrases where they truthfully apply; never fabricate: {metrics}"
        )
    extra = str(a.get("additional") or "").strip()
    if extra:
        lines.append(f"- Anything else to reflect (steers wording choice only): {extra}")

    if not lines:
        return ""
    return (
        "\n\nCANDIDATE STEER (these guide WHICH existing wording to rewrite and which "
        "keywords to weave in — they do NOT add content):\n"
        + "\n".join(lines)
        + "\nHARD RULE: rewrite ONLY text already present in the resume above. Do NOT "
        "invent skills, employers, metrics, or experience, and do NOT output any edit "
        "whose original_text is not an exact substring of the resume. If the steer asks "
        "for something not in the resume, partially honor it through wording choice only."
    )


def compute_keyword_edits(
    resume_text: str,
    jd_text: str,
    *,
    answers: dict[str, Any] | None = None,
    settings: Settings | None = None,
    client: Any | None = None,
) -> list[dict[str, str]]:
    """Ask the model for minimal {original_text → replacement_text} swaps.
    `answers` (the 6 customization fields) steer WHICH existing phrases/keywords
    to favour — never to add content. Demo mode (no key) returns an empty list
    (no silent edits)."""
    settings = settings or get_settings()
    if not settings.has_anthropic_key:
        return []
    api = _build_client(settings, client)
    user_content = (
        f"RESUME TEXT:\n{resume_text}\n\nTARGET JOB:\n{jd_text}"
        + _keyword_addendum(answers)
        + "\n\nReturn the JSON edits."
    )
    messages: list[dict[str, Any]] = [{"role": "user", "content": user_content}]

    def _call() -> str:
        resp = api.messages.create(
            model=MODEL,
            max_tokens=1500,
            system=[{"type": "text", "text": _KEYWORD_SYSTEM}],
            messages=messages,
        )
        return _first_text(resp)

    # Prompt-based JSON (NOT structured-output/grammar — that 400s on complex
    # schemas). Parse, and on a syntax error retry ONCE with a corrective turn;
    # a second failure raises KeywordInjectionError for a clean user message.
    text = _call()
    try:
        obj = _extract_json_object(text)
    except (json.JSONDecodeError, ValueError) as e:
        keyword_parse_failures["first"] += 1
        log.warning("ats keyword-inject: first parse failed (%s) — retrying", e)
        messages.append({"role": "assistant", "content": text})
        messages.append({"role": "user", "content": _KEYWORD_JSON_FIX.format(error=e)})
        retry_text = _call()
        try:
            obj = _extract_json_object(retry_text)
        except (json.JSONDecodeError, ValueError) as e2:
            keyword_parse_failures["retry"] += 1
            log.warning("ats keyword-inject: retry parse also failed (%s)", e2)
            raise KeywordInjectionError(
                "Couldn't generate keyword edits — try again, or use the standard "
                "JD-paste flow instead."
            ) from e2
    out: list[dict[str, str]] = []
    for e in obj.get("edits", []):
        o, r = str(e.get("original_text", "")).strip(), str(e.get("replacement_text", "")).strip()
        if o and r and o != r:
            out.append(
                {"original_text": o, "replacement_text": r, "reason": str(e.get("reason", ""))}
            )
    # Steer-only is enforced at apply time: apply_docx_edits APPLIES only edits
    # whose original is already in the resume; anything absent is reported in
    # `skipped` (surfaced read-only as "consider adding to your base resume"),
    # never inserted. The prompt's HARD RULE steers the model the same way.
    return out


def extract_docx_text(docx_bytes: bytes) -> str:
    """Plain-text extraction of a DOCX (python-docx), for the keyword model."""
    from docx import Document  # noqa: PLC0415

    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def apply_docx_edits(
    docx_bytes: bytes, edits: list[dict[str, str]]
) -> tuple[bytes, list[dict[str, str]], list[dict[str, str]]]:
    """Apply {original_text → replacement_text} swaps to a DOCX in place,
    preserving formatting. Matching is PARAGRAPH-LEVEL: an edit applies even
    when `original_text` spans multiple runs (Word splits text into runs at
    arbitrary boundaries), via run-splicing — the full replacement goes into the
    FIRST overlapped run and the overlapped remainder is blanked from the other
    runs. Every other run is left untouched, so formatting (bold/size/colour/
    etc.) is preserved. An edit is skipped ONLY when its text isn't present in
    any single paragraph at all.

    Returns (new_docx_bytes, applied, skipped)."""
    from docx import Document  # noqa: PLC0415

    doc = Document(io.BytesIO(docx_bytes))
    applied: list[dict[str, str]] = []
    skipped: list[dict[str, str]] = []

    # All runs across body paragraphs + table cells.
    def _all_paragraphs():  # noqa: ANN202
        yield from doc.paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    paragraphs = list(_all_paragraphs())
    for edit in edits:
        if _apply_one_edit(paragraphs, edit["original_text"], edit["replacement_text"]):
            applied.append(edit)
        else:
            log.info(
                "ats keyword-inject: skipped edit (not found in any paragraph): %r",
                edit["original_text"][:60],
            )
            skipped.append(edit)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), applied, skipped


def _apply_one_edit(paragraphs: list, original: str, replacement: str) -> bool:
    """Replace the FIRST occurrence of `original` in the FIRST paragraph that
    contains it (across run boundaries), preserving each run's OWN formatting.

    The replacement is DISTRIBUTED across the overlapped runs in proportion to
    each run's share of the original span — so when a match crosses runs of
    differing style (e.g. a bold label + a normal value), the bold portion stays
    bold and the normal portion stays normal, instead of the whole replacement
    inheriting the first run's style. Untouched runs are unchanged; no new runs,
    no net-new lines. Returns True if applied."""
    if not original:
        return False
    for para in paragraphs:
        runs = para.runs
        if not runs:
            continue
        # Snapshot original run texts BEFORE mutating, so the offset math is
        # immune to length changes we make along the way.
        texts = [r.text for r in runs]
        full = "".join(texts)
        idx = full.find(original)
        if idx < 0:
            continue
        end = idx + len(original)
        orig_len = len(original)

        # First pass: collect the overlapped runs (run, head, tail, frac) where
        # `frac` is the run's share of the original span. head/tail are the
        # parts of the run OUTSIDE the match, which we always keep verbatim.
        spans = []
        pos = 0
        for run, t in zip(runs, texts, strict=True):
            r_start, r_end = pos, pos + len(t)
            pos = r_end
            ov_start, ov_end = max(idx, r_start), min(end, r_end)
            if ov_start >= ov_end:
                continue
            ls, le = ov_start - r_start, ov_end - r_start
            frac = (ov_end - ov_start) / orig_len if orig_len else 0.0
            spans.append({"run": run, "head": t[:ls], "tail": t[le:], "frac": frac})

        # Second pass: slice `replacement` across the overlapped runs by `frac`,
        # so each run carries the proportional chunk in ITS own formatting. The
        # LAST overlapped run takes the remainder (no characters dropped to
        # rounding). Single-overlap case → the whole replacement in that run.
        n = len(spans)
        cursor = 0
        for i, sp in enumerate(spans):
            if i == n - 1:
                chunk = replacement[cursor:]
            else:
                take = round(sp["frac"] * len(replacement))
                chunk = replacement[cursor : cursor + take]
                cursor += take
            sp["run"].text = sp["head"] + chunk + sp["tail"]
        return True
    return False
