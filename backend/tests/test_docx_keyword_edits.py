"""apply_docx_edits: cross-run, paragraph-level in-place keyword swaps.

Word splits a paragraph's text into runs at arbitrary boundaries, so a
multi-word `original_text` usually spans several runs. The splicer must apply
it anyway (putting the replacement in the first overlapped run and blanking the
overlapped remainder of the others) while preserving every other run's
formatting — and skip ONLY when the text genuinely isn't in any paragraph.
"""

from __future__ import annotations

import io

from docx import Document

from app.services import ats


def _docx(paras: list[list[tuple[str, bool | None]]]) -> bytes:
    """Build a DOCX. Each paragraph is a list of (text, bold) run tuples — i.e.
    a fixture with EXPLICIT run splits (unlike add_paragraph's single run)."""
    doc = Document()
    for runs in paras:
        p = doc.add_paragraph()
        for text, bold in runs:
            r = p.add_run(text)
            r.bold = bold
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_cross_run_edit_applied_and_formatting_preserved():
    # "Python and Kafka" spans runs 1–3 (with bold on "Python " and "Kafka ").
    blob = _docx(
        [
            [
                ("I use ", None),
                ("Python ", True),
                ("and ", None),
                ("Kafka ", True),
                ("daily.", None),
            ],
            [("Untouched bullet.", None)],
        ]
    )
    edited, applied, skipped = ats.apply_docx_edits(
        blob, [{"original_text": "Python and Kafka", "replacement_text": "Python, Kafka, Spark"}]
    )
    assert len(applied) == 1
    assert skipped == []
    d = Document(io.BytesIO(edited))
    # The swap landed across the run boundaries.
    assert d.paragraphs[0].text == "I use Python, Kafka, Spark daily."
    # Formatting preserved: the replacement inherits the FIRST overlapped run's
    # bold, and bold runs still exist in the paragraph.
    assert any(r.bold for r in d.paragraphs[0].runs if r.text)


def test_genuinely_absent_string_is_skipped():
    blob = _docx([[("Hello ", None), ("world.", True)]])
    edited, applied, skipped = ats.apply_docx_edits(
        blob, [{"original_text": "Goodbye", "replacement_text": "Hi"}]
    )
    assert applied == []
    assert len(skipped) == 1
    assert Document(io.BytesIO(edited)).paragraphs[0].text == "Hello world."


def test_untouched_paragraphs_are_byte_identical():
    blob = _docx(
        [
            [("Edit ", None), ("me ", True), ("now.", None)],
            [("Leave ", None), ("me ", True), ("alone.", None)],
        ]
    )
    edited, applied, _skipped = ats.apply_docx_edits(
        blob, [{"original_text": "Edit me", "replacement_text": "Changed it"}]
    )
    assert len(applied) == 1
    d0 = Document(io.BytesIO(blob))
    d1 = Document(io.BytesIO(edited))
    assert d1.paragraphs[0].text == "Changed it now."  # edited (spanned 2 runs)
    # The untouched paragraph serialises byte-for-byte identically.
    assert d0.paragraphs[1]._p.xml == d1.paragraphs[1]._p.xml


def test_single_run_edit_still_works():
    # Backward-compat: an original fully inside one run behaves as before.
    blob = _docx([[("Python developer.", None)]])
    edited, applied, _skipped = ats.apply_docx_edits(
        blob, [{"original_text": "Python", "replacement_text": "Python and Kafka"}]
    )
    assert len(applied) == 1
    assert Document(io.BytesIO(edited)).paragraphs[0].text == "Python and Kafka developer."


def test_apply_is_insertion_free():
    # Option B never inserts: apply only rewrites EXISTING runs. An absent
    # original is reported skipped (a suggestion), never added as a new line.
    blob = _docx([[("Alpha line.", None)], [("Beta line.", None)]])
    edited, applied, skipped = ats.apply_docx_edits(
        blob,
        [
            {"original_text": "Alpha", "replacement_text": "Alpha Plus"},
            {"original_text": "Gamma", "replacement_text": "Gamma X"},  # absent
        ],
    )
    d0 = Document(io.BytesIO(blob))
    d1 = Document(io.BytesIO(edited))
    # No new paragraphs/lines/sections.
    assert len(d1.paragraphs) == len(d0.paragraphs) == 2
    assert len(applied) == 1 and len(skipped) == 1  # absent → skipped, not inserted
    assert d1.paragraphs[0].text == "Alpha Plus line."
    assert d1.paragraphs[1].text == "Beta line."  # untouched
