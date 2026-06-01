"""Option-B customize + selective-revert + custom filename.

(a) compute_keyword_edits folds the 6 customization answers (incl. free-text
    "additional") into the prompt as a STEER, and still returns swaps that only
    touch existing text (original_text verbatim in the resume).
(b) downloading with accepted=[subset] rebuilds the DOCX from ONLY those edits;
    untouched paragraphs are byte-identical.
(c) filename sanitization strips path separators / quotes and forces a single
    `.docx`; blank → default.
"""

from __future__ import annotations

import io

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.api.ats import _sanitize_docx_filename
from app.config import Settings
from app.database import Base
from app.models.tailor_run import TailorRun
from app.models.user import User
from app.services import ats


def _key_settings() -> Settings:
    return Settings(DATABASE_URL="x", ADMIN_TOKEN="t", ANTHROPIC_API_KEY="sk-x")


def _docx_bytes(paragraphs: list[str]) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── (a) answers fold into the keyword prompt ─────────────────────────────────


class _Capture:
    """Fake Anthropic client: records the user prompt, returns a scripted reply."""

    _DEFAULT = '{"edits":[{"original_text":"Python","replacement_text":"Python, Kafka"}]}'

    def __init__(self, reply: str | None = None) -> None:
        self.messages = self
        self.user = ""
        self._reply = reply or self._DEFAULT

    def create(self, **kw):  # noqa: ANN003
        self.user = kw["messages"][0]["content"]
        block = type("B", (), {"type": "text", "text": self._reply})()
        return type("M", (), {"content": [block]})()


def test_compute_keyword_edits_folds_option_b_answers_and_only_touches_existing():
    # Model returns one PRESENT-original edit ("Python") + one ABSENT one ("Java").
    cap = _Capture(
        '{"edits":['
        '{"original_text":"Python","replacement_text":"Python, Kafka"},'
        '{"original_text":"Java","replacement_text":"Java, Scala"}]}'
    )
    resume_text = "I use Python daily."
    out = ats.compute_keyword_edits(
        resume_text,
        "Need Kafka.",
        answers={
            "missing_experience": "shipped a Kafka pipeline at scale",
            "skills": ["Kafka"],
            "roles": ["Engineer at Acme"],
            "metrics": "cut latency 40%",
            "additional": "lean into streaming systems",
        },
        settings=_key_settings(),
        client=cap,
    )
    # The 5-question steer (incl. the free-text fields) folds into the prompt,
    # with the HARD RULE.
    u = cap.user
    assert "CANDIDATE STEER" in u and "HARD RULE" in u
    assert "shipped a Kafka pipeline at scale" in u  # #1 missing_experience
    assert "Kafka" in u  # #2 skills-gap
    assert "Engineer at Acme" in u  # #3 most-relevant roles
    assert "cut latency 40%" in u  # #4 metrics
    assert "lean into streaming systems" in u  # #5 free text

    # Steer-only: APPLY touches ONLY existing text — the present edit lands, the
    # absent one is skipped (a suggestion), never inserted.
    blob = _docx_bytes(["I use Python daily."])
    _, applied, skipped = ats.apply_docx_edits(blob, out)
    assert [e["original_text"] for e in applied] == ["Python"]
    assert [e["original_text"] for e in skipped] == ["Java"]
    assert all(e["original_text"] in resume_text for e in applied)
    assert out[0]["original_text"] in resume_text


def test_compute_keyword_edits_no_answers_no_steer_block():
    cap = _Capture()
    ats.compute_keyword_edits("resume", "jd", settings=_key_settings(), client=cap)
    assert "CANDIDATE STEER" not in cap.user


# ── (b)/(c) download endpoint ────────────────────────────────────────────────


@pytest.fixture
def api(tmp_path):
    from app import config as config_module
    from app.api.auth import get_current_user
    from app.config import get_settings
    from app.database import get_db
    from app.main import app

    engine = create_engine(
        f"sqlite+pysqlite:///{tmp_path / 'cd.db'}",
        future=True,
        connect_args={"check_same_thread": False},
    )
    Base.metadata.create_all(engine)
    Session = sessionmaker(bind=engine, future=True)
    with Session() as s:
        s.add(User(id=1, google_subject_id="g", email="t@e.com", name="T"))
        s.commit()
    user = type("U", (), {"id": 1, "email": "t@e.com", "name": "T"})()

    def override_db():
        with Session() as s:
            yield s

    app.dependency_overrides[get_db] = override_db
    app.dependency_overrides[get_settings] = lambda: Settings(
        DATABASE_URL="x", ADMIN_TOKEN="t", ANTHROPIC_API_KEY=""
    )
    app.dependency_overrides[get_current_user] = lambda: user
    config_module.get_settings.cache_clear()
    try:
        yield TestClient(app), Session
    finally:
        app.dependency_overrides.clear()
        config_module.get_settings.cache_clear()


def _park_run(Session, run_id: str, blob: bytes, applied: list[dict]) -> None:
    with Session() as s:
        s.add(
            TailorRun(
                run_id=run_id,
                user_id=1,
                option_type="upload_docx",
                uploaded_filename="me.docx",
                uploaded_docx_blob=blob,
                status="done",
                result_json={"kind": "docx_keyword_inject", "applied": applied, "skipped": []},
            )
        )
        s.commit()


def test_download_accepted_subset_rebuilds_only_those_edits(api):
    client, Session = api
    blob = _docx_bytes(["Alpha line.", "Beta line."])
    _park_run(
        Session,
        "r1",
        blob,
        [
            {"original_text": "Alpha", "replacement_text": "Alpha+"},
            {"original_text": "Beta", "replacement_text": "Beta+"},
        ],
    )
    # Accept only edit 0.
    res = client.post("/api/ats/runs/r1/download-docx", json={"accepted": [0]})
    assert res.status_code == 200
    d = Document(io.BytesIO(res.content))
    assert d.paragraphs[0].text == "Alpha+ line."  # edit 0 applied
    assert d.paragraphs[1].text == "Beta line."  # edit 1 NOT applied (rejected)
    # The rejected/untouched paragraph is byte-identical to the original.
    d0 = Document(io.BytesIO(blob))
    assert d0.paragraphs[1]._p.xml == d.paragraphs[1]._p.xml


def test_filename_sanitization_unit():
    f = _sanitize_docx_filename
    assert f("../../etc/passwd", default="d") == "passwd.docx"  # path stripped
    assert f('my "resume".docx', default="d") == "my resume.docx"  # quotes + ext
    assert f("report", default="d") == "report.docx"
    assert f("report.DOCX", default="d") == "report.docx"  # case-insensitive ext
    assert f("", default="fallback") == "fallback.docx"  # blank → default
    assert f(None, default="fallback") == "fallback.docx"
    long = f("a" * 300, default="d")
    assert long.endswith(".docx") and len(long) <= 133


def test_download_uses_sanitized_filename(api):
    client, Session = api
    _park_run(Session, "r2", _docx_bytes(["Alpha line."]), [])
    res = client.post("/api/ats/runs/r2/download-docx", json={"filename": "My Final/Resume"})
    assert res.status_code == 200
    assert 'filename="Resume.docx"' in res.headers["content-disposition"]


def test_download_blank_filename_falls_back_to_default(api):
    client, Session = api
    _park_run(Session, "r3", _docx_bytes(["Alpha line."]), [])
    res = client.post("/api/ats/runs/r3/download-docx", json={})
    assert res.status_code == 200
    # Default derived from the uploaded name: me.docx → me-tailored.docx
    assert 'filename="me-tailored.docx"' in res.headers["content-disposition"]
