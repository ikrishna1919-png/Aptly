"""PDF-direct parsing + DOCX-completeness regression tests.

Two product changes under test:

1. The PDF upload path no longer extracts text with pdfplumber and
   feeds the text to Claude — it sends the PDF bytes to the
   Anthropic API as a `document` content block. The worker uses a
   different background-launch function (`start_background_parse_pdf`)
   and a different LLM helper (`_llm_extract_structural_pdf`).

2. The DOCX text extractor now walks every text-bearing surface
   (paragraphs with list-marker preservation, table cells,
   headers/footers, text boxes) so bullets and table-housed contact
   rows survive to the parser.

Both halves are exercised here without hitting the real Anthropic
API — the SDK is monkey-patched so the assertions cover the
request-shape and the routing, not the model output.
"""

from __future__ import annotations

import io
from types import SimpleNamespace

import pytest
from sqlalchemy import create_engine, select
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app.config import Settings
from app.database import Base
from app.models.parse_run import PARSE_STATUS_RUNNING, PARSE_STATUS_SUCCESS, ParseRun
from app.services import profile_parser as parser_module
from app.services.resume_extractor import _extract_docx


@pytest.fixture
def factories():
    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        future=True,
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    return sessionmaker(bind=engine, future=True)


# ─── DOCX completeness ─────────────────────────────────────────────────────


def _docx_bytes(builder) -> bytes:
    """Run `builder(doc)` then return the DOCX file's bytes."""
    from docx import Document

    doc = Document()
    builder(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxCompleteness:
    def test_list_paragraphs_are_prefixed_for_bullet_recognition(self):
        """A DOCX bulleted list shows up as paragraphs with `style.name
        == "List Bullet"`. The new extractor prefixes those lines with
        `- ` so the downstream parser recognises them as bullets even
        though python-docx's `para.text` strips the visual glyph."""

        def build(doc):
            doc.add_paragraph("Jordan Singh")
            doc.add_paragraph("Senior Engineer · Acme")
            doc.add_paragraph("Cut p95 latency 480ms→110ms", style="List Bullet")
            doc.add_paragraph("Adopted by 6 teams", style="List Bullet")
            doc.add_paragraph("Education")

        text = _extract_docx(_docx_bytes(build))
        assert "- Cut p95 latency 480ms→110ms" in text
        assert "- Adopted by 6 teams" in text
        # Non-list paragraphs are NOT prefixed.
        assert "Jordan Singh" in text and "- Jordan Singh" not in text

    def test_table_cells_are_extracted_with_tab_separation(self):
        """Two-column resume templates park dates / titles in tables.
        Cells survive as tab-separated rows."""

        def build(doc):
            doc.add_paragraph("Header")
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Senior Engineer, Acme"
            table.rows[0].cells[1].text = "2022 – Present"
            table.rows[1].cells[0].text = "Engineer, Beta"
            table.rows[1].cells[1].text = "2019 – 2022"

        text = _extract_docx(_docx_bytes(build))
        # Every row joins with a tab between cells.
        assert "Senior Engineer, Acme\t2022 – Present" in text
        assert "Engineer, Beta\t2019 – 2022" in text

    def test_header_and_footer_content_is_captured(self):
        """Enterprise templates put contact info in the header. The
        extractor walks every section's header + footer so those
        rows survive."""

        def build(doc):
            section = doc.sections[0]
            section.header.paragraphs[0].text = "jordan@example.com · 555-123-4567"
            section.footer.paragraphs[0].text = "Page 1 of 1"
            doc.add_paragraph("Body paragraph")

        text = _extract_docx(_docx_bytes(build))
        assert "jordan@example.com · 555-123-4567" in text
        assert "Page 1 of 1" in text
        assert "Body paragraph" in text

    def test_empty_paragraphs_preserved_as_blank_lines(self):
        """The section segmenter splits entries on blank lines.
        Empty DOCX paragraphs must round-trip as blank lines so a
        DOCX-uploaded resume sections the same way a pasted one
        does."""

        def build(doc):
            doc.add_paragraph("Experience")
            doc.add_paragraph("")
            doc.add_paragraph("Senior Engineer")

        text = _extract_docx(_docx_bytes(build))
        assert "Experience\n\nSenior Engineer" in text

    def test_empty_docx_raises_empty_extraction(self):
        from app.services.resume_extractor import EmptyExtractionError, extract_text

        # A DOCX with one truly-empty paragraph normalises to "".
        def build(doc):
            doc.add_paragraph("")

        with pytest.raises(EmptyExtractionError, match="paste"):
            extract_text("resume.docx", _docx_bytes(build))


# ─── PDF-direct path ───────────────────────────────────────────────────────


def _seed_run(Session, run_id: str = "run-pdf") -> None:
    with Session() as s:
        s.add(ParseRun(run_id=run_id, status=PARSE_STATUS_RUNNING, user_id=None))
        s.commit()


def _row(Session, run_id: str) -> ParseRun:
    with Session() as s:
        return s.execute(select(ParseRun).where(ParseRun.run_id == run_id)).scalar_one()


class _MockAnthropicClient:
    """Captures the call shape so we can assert the request that went
    to the API; returns a canned structured response so the worker
    completes happily."""

    def __init__(self, payload: dict) -> None:
        self.calls: list[dict] = []
        self.payload = payload
        self.messages = SimpleNamespace(create=self._create)

    def _create(self, **kwargs):
        self.calls.append(kwargs)
        import json

        text_block = SimpleNamespace(type="text", text=json.dumps(self.payload))
        return SimpleNamespace(content=[text_block])


class TestPdfDirectPath:
    def test_pdf_worker_sends_document_content_block(self, factories, monkeypatch):
        """The PDF worker MUST send the PDF as a `document` content
        block, not as extracted text. This is the whole point of the
        new code path — Claude's document understanding handles
        bullets / multi-column / tables that pdfplumber loses."""
        payload = {
            "name": "Jordan Singh",
            "experience": [
                {
                    "company": "Acme",
                    "title": "Senior Engineer",
                    "location": "New York, NY",
                    "start_date": "Jan 2022",
                    "end_date": "Present",
                    "description_bullets": ["Built X.", "Shipped Y."],
                }
            ],
            "education": [],
            "skills": ["Python", "Go"],
            "projects": [],
            "achievements": [],
            "section_order": [],
        }
        mock = _MockAnthropicClient(payload)
        monkeypatch.setattr(parser_module, "_build_client", lambda s, c: mock)
        monkeypatch.setattr(parser_module, "SessionLocal", factories)
        monkeypatch.setattr(parser_module, "_launch_worker", lambda t, a: t(*a))

        _seed_run(factories, "pdf-1")
        settings = Settings(
            DATABASE_URL="sqlite+pysqlite:///:memory:",
            ADMIN_TOKEN="t",
            ANTHROPIC_API_KEY="sk-test",
        )
        # Tiny pseudo-PDF — pdfplumber will fail to parse it, but the
        # worker logs + continues; the LLM call (mocked) is what
        # delivers the result.
        parser_module._execute_parse_run_pdf("pdf-1", b"%PDF-1.4\nfake", settings)

        run = _row(factories, "pdf-1")
        assert run.status == PARSE_STATUS_SUCCESS
        assert run.profile["name"] == "Jordan Singh"
        assert run.profile["experience"][0]["company"] == "Acme"

        # Inspect the call shape — exactly one Anthropic call, and
        # the message content carries a `document` block with the
        # base64-encoded PDF AND a follow-up text instruction.
        assert len(mock.calls) == 1
        call = mock.calls[0]
        content = call["messages"][0]["content"]
        assert isinstance(content, list)
        doc_blocks = [b for b in content if b.get("type") == "document"]
        text_blocks = [b for b in content if b.get("type") == "text"]
        assert len(doc_blocks) == 1, f"expected one document block, got: {content!r}"
        assert doc_blocks[0]["source"]["media_type"] == "application/pdf"
        assert doc_blocks[0]["source"]["type"] == "base64"
        # The base64 payload decodes back to the original PDF bytes.
        import base64

        decoded = base64.standard_b64decode(doc_blocks[0]["source"]["data"])
        assert decoded == b"%PDF-1.4\nfake"
        # And the text instruction is present so the model knows
        # what to do.
        assert any("attached PDF" in (b.get("text") or "") for b in text_blocks)

    def test_pdf_llm_failure_falls_back_to_regex_pass(self, factories, monkeypatch):
        """If the Anthropic call raises, the worker still writes a
        terminal `success` row — using whatever the regex contact-
        field pass managed to extract from pdfplumber. The error is
        logged but doesn't propagate."""

        class _Raises:
            def create(self, **_):
                raise RuntimeError("anthropic 500")

        monkeypatch.setattr(
            parser_module,
            "_build_client",
            lambda s, c: SimpleNamespace(messages=_Raises()),
        )
        monkeypatch.setattr(parser_module, "SessionLocal", factories)
        monkeypatch.setattr(parser_module, "_launch_worker", lambda t, a: t(*a))

        _seed_run(factories, "pdf-fail")
        settings = Settings(
            DATABASE_URL="sqlite+pysqlite:///:memory:",
            ADMIN_TOKEN="t",
            ANTHROPIC_API_KEY="sk-test",
        )
        parser_module._execute_parse_run_pdf("pdf-fail", b"%PDF-1.4\nfake", settings)
        run = _row(factories, "pdf-fail")
        # Worker writes a terminal status even when the LLM fails —
        # never sits at running.
        assert run.status == PARSE_STATUS_SUCCESS
        # No profile content because pdfplumber couldn't read the
        # fake PDF either; regex returned an empty profile.
        assert run.profile["name"] == ""

    def test_start_background_parse_pdf_creates_run_row(self, factories, monkeypatch):
        """Sanity: the PDF launcher actually creates a ParseRun row
        with the right user_id + initial running status."""
        monkeypatch.setattr(parser_module, "SessionLocal", factories)
        monkeypatch.setattr(parser_module, "_launch_worker", lambda *_a, **_k: None)

        settings = Settings(DATABASE_URL="sqlite+pysqlite:///:memory:", ADMIN_TOKEN="t")
        run_id = parser_module.start_background_parse_pdf(
            b"%PDF-1.4\nfake", user_id=42, settings=settings
        )
        with factories() as s:
            row = s.execute(select(ParseRun).where(ParseRun.run_id == run_id)).scalar_one()
        assert row.user_id == 42
        assert row.status == "running"


class TestUploadEndpointPdfRouting:
    """The HTTP edge of the upload endpoint MUST hand a PDF to
    `start_background_parse_pdf`, not the text-input launcher."""

    def test_pdf_upload_goes_through_pdf_launcher(self, monkeypatch):
        from fastapi.testclient import TestClient

        import app.api.profile as profile_api
        from app import config as config_module
        from app.api.auth import get_current_user
        from app.config import get_settings as get_settings_dep
        from app.database import get_db
        from app.main import app
        from app.models.user import User

        engine = create_engine(
            "sqlite+pysqlite:///:memory:",
            future=True,
            connect_args={"check_same_thread": False},
            poolclass=StaticPool,
        )
        Base.metadata.create_all(engine)
        Session = sessionmaker(bind=engine, future=True)
        with Session() as s:
            user = User(google_subject_id="g", email="u@x.com", name="U")
            s.add(user)
            s.commit()
            s.refresh(user)
            s.expunge(user)

        def override_db():
            with Session() as s:
                yield s

        captured: dict[str, object] = {}

        def fake_pdf(data, *, user_id, settings):
            captured["pdf"] = data
            captured["user_id"] = user_id
            return "fake-run-id"

        def fake_text(*_a, **_k):
            captured["text"] = "should-not-be-called"
            return "should-not"

        monkeypatch.setattr(profile_api, "start_background_parse_pdf", fake_pdf)
        monkeypatch.setattr(profile_api, "start_background_parse", fake_text)

        settings = Settings(DATABASE_URL="sqlite+pysqlite:///:memory:", ADMIN_TOKEN="t")
        app.dependency_overrides[get_db] = override_db
        app.dependency_overrides[get_settings_dep] = lambda: settings
        app.dependency_overrides[get_current_user] = lambda: user
        config_module.get_settings.cache_clear()
        try:
            files = {"file": ("resume.pdf", b"%PDF-1.4\nhello", "application/pdf")}
            res = TestClient(app).post("/api/profile/parse/upload", files=files)
            assert res.status_code == 202
        finally:
            app.dependency_overrides.clear()
            config_module.get_settings.cache_clear()

        assert captured["pdf"] == b"%PDF-1.4\nhello"
        assert captured["user_id"] == user.id
        # The text launcher is NOT called for a PDF upload — the
        # direct-to-LLM path bypasses pdfplumber-then-text entirely.
        assert "text" not in captured

    def test_obviously_corrupt_pdf_400s_at_edge(self, monkeypatch):
        """Magic-byte sanity check stops a non-PDF before we burn an
        Anthropic call."""
        from fastapi.testclient import TestClient

        from app import config as config_module
        from app.api.auth import get_current_user
        from app.config import get_settings as get_settings_dep
        from app.database import get_db
        from app.main import app
        from app.models.user import User

        engine = create_engine(
            "sqlite+pysqlite:///:memory:",
            future=True,
            connect_args={"check_same_thread": False},
            poolclass=StaticPool,
        )
        Base.metadata.create_all(engine)
        Session = sessionmaker(bind=engine, future=True)
        with Session() as s:
            user = User(google_subject_id="g", email="u@x.com", name="U")
            s.add(user)
            s.commit()
            s.refresh(user)
            s.expunge(user)

        def override_db():
            with Session() as s:
                yield s

        settings = Settings(DATABASE_URL="sqlite+pysqlite:///:memory:", ADMIN_TOKEN="t")
        app.dependency_overrides[get_db] = override_db
        app.dependency_overrides[get_settings_dep] = lambda: settings
        app.dependency_overrides[get_current_user] = lambda: user
        config_module.get_settings.cache_clear()
        try:
            files = {"file": ("resume.pdf", b"not a pdf", "application/pdf")}
            res = TestClient(app).post("/api/profile/parse/upload", files=files)
            assert res.status_code == 400
            assert "PDF" in res.json()["detail"]
        finally:
            app.dependency_overrides.clear()
            config_module.get_settings.cache_clear()
