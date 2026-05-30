"""DOCX + PDF renderer tests — covers the ATS rendering contract:

- Single column; no tables, columns, text boxes, or images.
- Closed-list section headings.
- Visual mode: right-aligned date TAB STOP + heading rules.
- Plain mode: no rules; dates inline after the company line via " | ".
- Both modes carry identical wording (only styling differs).
- No en/em dashes or decorative bullet glyphs anywhere.
"""

from __future__ import annotations

import io
import re
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.docx_export import render_docx
from app.services.pdf_export import count_pages, render_pdf
from app.services.tailor import (
    Contact,
    EducationEntry,
    ExperienceEntry,
    ResumeMeta,
    SkillGroup,
    TailoredResume,
)


def _resume(**overrides) -> TailoredResume:
    defaults: dict = {
        "meta": ResumeMeta(mode="visual"),
        "contact": Contact(
            name="Alex Rivera",
            headline="Senior Software Engineer",
            location="San Francisco, CA",
            email="alex@example.com",
            phone="+1 (555) 123-4567",
        ),
        "summary": "Senior backend engineer with a strong Python and AWS background.",
        "skills": [
            SkillGroup(category="Languages", items=["Python", "TypeScript"]),
            SkillGroup(category="Cloud", items=["AWS", "Kubernetes"]),
        ],
        "experience": [
            ExperienceEntry(
                title="Senior Software Engineer",
                company="Forge Labs",
                location="San Francisco, CA",
                start_date="Feb 2023",
                end_date="Present",
                bullets=[
                    "Led billing migration to event-driven Kafka; cut p95 latency to 110ms.",
                    "Designed a feature-flag platform (FastAPI, Postgres) used by 6 teams.",
                ],
            ),
            ExperienceEntry(
                title="Software Engineer",
                company="Northwind Analytics",
                location="Remote",
                start_date="Jun 2020",
                end_date="Jan 2023",
                bullets=[
                    "Built a data ingestion pipeline (Airflow, Snowflake) for 4B events per day.",
                ],
            ),
        ],
        "education": [
            EducationEntry(
                degree="B.S. Computer Science",
                institution="Carnegie Mellon University",
                location="Pittsburgh, PA",
                graduation_date="May 2018",
            )
        ],
    }
    defaults.update(overrides)
    return TailoredResume(**defaults)


def _xml_blob(buf: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(buf)) as zf:
        return zf.read("word/document.xml").decode("utf-8")


def _docx_text(buf: bytes) -> str:
    return "\n".join(p.text for p in Document(io.BytesIO(buf)).paragraphs)


# ── No tables / columns / text boxes / images ──────────────────────────────


def test_no_tables_columns_or_text_boxes():
    for mode in ("visual", "plain"):
        xml = _xml_blob(render_docx(_resume(), mode=mode))
        assert "<w:tbl>" not in xml and "<w:tbl " not in xml
        cols_match = re.search(r'<w:cols[^>]*w:num="(\d+)"', xml)
        if cols_match:
            assert int(cols_match.group(1)) == 1, "document must be single-column"
        assert "txbxContent" not in xml
        assert "<w:drawing" not in xml and "<pic:pic" not in xml


# ── Visual mode: right tab stop + heading rules ────────────────────────────


def test_visual_mode_uses_right_tab_stop_for_dates():
    xml = _xml_blob(render_docx(_resume(), mode="visual"))
    right_tab_stops = [
        m
        for m in re.findall(r"<w:tab\b[^>]*/>", xml)
        if 'w:val="right"' in m and re.search(r'w:pos="\d+"', m)
    ]
    # One per experience + education entry (2 + 1 = 3 here).
    assert (
        len(right_tab_stops) >= 3
    ), f"expected right-aligned tab stops; got {len(right_tab_stops)}"
    # Inline `\t` between the company and the date run.
    assert re.findall(r"<w:tab\s*/>", xml)
    # Visual mode draws heading rules (paragraph bottom borders).
    assert "w:pBdr" in xml


def test_plain_mode_has_no_rules_and_inline_dates():
    buf = render_docx(_resume(), mode="plain")
    xml = _xml_blob(buf)
    # No heading rules, no right-aligned tab stops in plain mode.
    assert "w:pBdr" not in xml
    assert not [m for m in re.findall(r"<w:tab\b[^>]*/>", xml) if 'w:val="right"' in m]
    # Dates inline after the company/location, joined by " | ".
    text = _docx_text(buf)
    assert "Forge Labs, San Francisco, CA | Feb 2023 to Present" in text


# ── Headings (closed list) ─────────────────────────────────────────────────


def test_closed_list_headings_present():
    text = _docx_text(render_docx(_resume(), mode="visual"))
    for heading in ("Professional Summary", "Skills", "Experience", "Education"):
        assert heading in text
    # Skills render as labeled categories.
    assert "Languages: Python, TypeScript" in text


# ── Identical wording across modes ─────────────────────────────────────────


def test_modes_share_identical_words():
    """Visual and plain must contain the same words — only the date join
    differs ("\\t" vs " | "), which we normalise before comparing."""

    def words(buf: bytes) -> list[str]:
        norm = _docx_text(buf).replace("\t", " ").replace(" | ", " ")
        return re.findall(r"\S+", norm)

    assert words(render_docx(_resume(), mode="visual")) == words(
        render_docx(_resume(), mode="plain")
    )


# ── No disallowed characters survive to the rendered DOCX ──────────────────


def test_no_dashes_or_bullet_glyphs_in_docx():
    text = _docx_text(render_docx(_resume(), mode="visual"))
    for bad in ("–", "—", "•", "‘", "’", "“", "”"):
        assert bad not in text


# ── PDF renderer ───────────────────────────────────────────────────────────


def test_pdf_renders_both_modes():
    for mode in ("visual", "plain"):
        pdf = render_pdf(_resume(), mode=mode)
        assert pdf[:4] == b"%PDF"
        assert len(pdf) > 1000
        assert count_pages(_resume(), mode=mode) in (1, 2)


# ── Header alignment (additive; orthogonal to mode) ─────────────────────────

_DOCX_EXPECT = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def test_docx_header_alignment_applies_to_header_only():
    """Name (header) follows header_alignment; summary + experience lines stay
    left-aligned regardless."""
    for alignment in ("left", "center", "right"):
        for mode in ("visual", "plain"):
            doc = Document(
                io.BytesIO(render_docx(_resume(), mode=mode, header_alignment=alignment))
            )
            paras = doc.paragraphs
            name_p = next(p for p in paras if p.text == "Alex Rivera")
            assert name_p.alignment == _DOCX_EXPECT[alignment], (alignment, mode)
            # Body text never follows the header alignment.
            summary_p = next(p for p in paras if p.text.startswith("Senior backend engineer"))
            assert summary_p.alignment in (None, WD_ALIGN_PARAGRAPH.LEFT)
            company_p = next(p for p in paras if "Forge Labs" in p.text)
            assert company_p.alignment in (None, WD_ALIGN_PARAGRAPH.LEFT)


def test_docx_default_header_alignment_is_center():
    doc = Document(io.BytesIO(render_docx(_resume())))
    name_p = next(p for p in doc.paragraphs if p.text == "Alex Rivera")
    assert name_p.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_pdf_renders_all_alignments():
    for alignment in ("left", "center", "right"):
        for mode in ("visual", "plain"):
            pdf = render_pdf(_resume(), mode=mode, header_alignment=alignment)
            assert pdf[:4] == b"%PDF"
            assert len(pdf) > 1000


def test_pdf_styles_align_header_not_body():
    """The header styles follow the choice; body + the entry meta line
    (shared 'contact' style) stay left so experience/education don't move."""
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

    from app.services.pdf_export import _styles

    want = {"left": TA_LEFT, "center": TA_CENTER, "right": TA_RIGHT}
    for alignment, expected in want.items():
        styles = _styles(alignment)
        assert styles["name"].alignment == expected
        assert styles["headline"].alignment == expected
        assert styles["header_contact"].alignment == expected
        # Body + the Entry meta line must NOT move.
        assert styles["body"].alignment == TA_LEFT
        assert styles["contact"].alignment == TA_LEFT
        assert styles["entry1"].alignment == TA_LEFT
