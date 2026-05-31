"""Tests for the /ats resume hub: formats, customization, DOCX keyword
injection, and the endpoints (demo mode, inline workers)."""

from __future__ import annotations

import io

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, select
from sqlalchemy.orm import sessionmaker

from app.config import Settings
from app.database import Base
from app.models.candidate import DEMO_SLUG, Candidate
from app.models.tailor_run import TailorRun
from app.services import ats, ats_runs
from app.services.ats import apply_docx_edits, build_customization_addendum
from app.services.docx_export import render_docx
from app.services.pdf_export import render_pdf
from app.services.resume_formats import VALID_FORMATS, resolve_format
from app.services.tailor import (
    Contact,
    ExperienceEntry,
    ResumeMeta,
    SkillGroup,
    TailoredResume,
)


def _no_key() -> Settings:
    return Settings(
        DATABASE_URL="sqlite+pysqlite:///:memory:", ADMIN_TOKEN="t", ANTHROPIC_API_KEY=""
    )


def _resume() -> TailoredResume:
    return TailoredResume(
        meta=ResumeMeta(),
        contact=Contact(name="Jane Doe", headline="Engineer", location="NYC", email="j@e.com"),
        summary="Backend engineer with Python.",
        skills=[SkillGroup(category="Languages", items=["Python"])],
        experience=[
            ExperienceEntry(
                title="Engineer",
                company="Acme",
                location="NYC",
                start_date="Jan 2022",
                end_date="Present",
                bullets=["Built Python services."],
            )
        ],
    )


# ─── Formats ──────────────────────────────────────────────────────────────────


class TestFormats:
    def test_resolve_presets(self):
        for name in ("modern", "classic", "minimal", "plain"):
            assert resolve_format(name).name == name
        assert resolve_format("nonsense").name == "modern"  # safe fallback

    def test_custom_overrides(self):
        spec = resolve_format(
            "custom",
            {"base": "classic", "accent_color": "teal", "font_family": "serif", "margins": "loose"},
        )
        assert spec.name == "custom" and spec.serif and spec.margins == 0.85
        assert spec.accent == (0x0F, 0x76, 0x6E)

    def test_all_formats_render_docx_and_pdf(self):
        r = _resume()
        for fmt in VALID_FORMATS:
            custom = {"base": "modern"} if fmt == "custom" else None
            d = render_docx(r, fmt=fmt, custom=custom)
            p = render_pdf(r, fmt=fmt, custom=custom)
            assert d[:2] == b"PK" and len(d) > 1500
            assert p[:4] == b"%PDF" and len(p) > 1000

    def test_formats_keep_58_contract_no_dashes(self):
        r = _resume()
        for fmt in ("modern", "classic", "minimal", "plain"):
            text = "\n".join(
                p.text for p in Document(io.BytesIO(render_docx(r, fmt=fmt))).paragraphs
            )
            for bad in ("–", "—", "•"):
                assert bad not in text


# ─── Customization ──────────────────────────────────────────────────────────


def test_customization_addendum():
    out = build_customization_addendum(
        {
            "length": "2",
            "tone": "formal",
            "emphasis": "leadership",
            "skills": ["Python", "Kafka"],
            "roles": ["Engineer at Acme"],
            "additional": "Avoid my 2019 internship.",
        }
    )
    assert "2 pages" in out
    assert "formal" in out.lower()
    assert "leadership" in out.lower()
    assert "Python, Kafka" in out
    assert "Engineer at Acme" in out
    assert "2019 internship" in out


def test_customization_defaults_to_one_page():
    assert "1 page" in build_customization_addendum({})


# ─── DOCX keyword injection ───────────────────────────────────────────────────


def _docx_with_runs(*paragraphs: list[str]) -> bytes:
    """Build a DOCX where each paragraph is a list of run texts."""
    doc = Document()
    for runs in paragraphs:
        p = doc.add_paragraph()
        for t in runs:
            p.add_run(t)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxKeywordInjection:
    def test_single_run_edit_applied_and_format_preserved(self):
        blob = _docx_with_runs(["Built Python data pipelines."])
        new, applied, skipped = apply_docx_edits(
            blob, [{"original_text": "Python", "replacement_text": "Python and Kafka"}]
        )
        assert len(applied) == 1 and not skipped
        text = "\n".join(p.text for p in Document(io.BytesIO(new)).paragraphs)
        assert "Python and Kafka data pipelines" in text

    def test_cross_run_edit_applied(self):
        # "Built pipelines" spans two separate runs → now APPLIED via run-
        # splicing (previously skipped). The whole replacement lands in the
        # first overlapped run; the rest of the span is blanked.
        blob = _docx_with_runs(["Built ", "pipelines"])
        new, applied, skipped = apply_docx_edits(
            blob,
            [{"original_text": "Built pipelines", "replacement_text": "Built Kafka pipelines"}],
        )
        assert len(applied) == 1 and not skipped
        text = "\n".join(p.text for p in Document(io.BytesIO(new)).paragraphs)
        assert text == "Built Kafka pipelines"

    def test_absent_edit_skipped(self):
        blob = _docx_with_runs(["Hello world."])
        _, applied, skipped = apply_docx_edits(
            blob, [{"original_text": "Goodbye", "replacement_text": "Hi"}]
        )
        assert not applied and len(skipped) == 1

    def test_compute_edits_demo_mode_returns_empty(self):
        assert ats.compute_keyword_edits("resume text", "jd text", settings=_no_key()) == []


# ─── Endpoints (demo mode, inline workers) ────────────────────────────────────


@pytest.fixture
def api(tmp_path, monkeypatch):
    from app import config as config_module
    from app.api.auth import get_current_user
    from app.config import get_settings
    from app.database import get_db
    from app.main import app
    from app.models.user import User

    engine = create_engine(
        f"sqlite+pysqlite:///{tmp_path/'ats.db'}",
        future=True,
        connect_args={"check_same_thread": False},
    )
    Base.metadata.create_all(engine)
    Session = sessionmaker(bind=engine, future=True)
    with Session() as s:
        u = User(google_subject_id="g", email="t@e.com", name="T")
        s.add(u)
        s.add(
            Candidate(
                slug=DEMO_SLUG,
                user_id=None,
                profile={
                    "name": "Jane Doe",
                    "email": "j@e.com",
                    "summary": "Backend eng.",
                    "skills": ["Python"],
                    "experience": [
                        {
                            "title": "Eng",
                            "company": "Acme",
                            "start": "2022-01",
                            "end": "Present",
                            "bullets": ["Built services."],
                        }
                    ],
                },
            )
        )
        s.commit()
        uid = u.id
    user = type("U", (), {"id": uid, "email": "t@e.com"})()
    settings = _no_key()

    monkeypatch.setattr(ats_runs, "SessionLocal", Session)
    monkeypatch.setattr(ats_runs, "_launch", lambda target, args: target(*args))

    def override_db():
        with Session() as s:
            yield s

    app.dependency_overrides[get_db] = override_db
    app.dependency_overrides[get_settings] = lambda: settings
    app.dependency_overrides[get_current_user] = lambda: user
    config_module.get_settings.cache_clear()
    try:
        yield TestClient(app), Session
    finally:
        app.dependency_overrides.clear()
        config_module.get_settings.cache_clear()


def test_generate_jd_paste_reaches_done(api):
    client, Session = api
    res = client.post(
        "/api/ats/generate",
        json={
            "option_type": "jd_paste",
            "jd_text": "Python and Kafka role. " * 20,
            "questions": {"length": "1", "tone": "confident"},
            "format": "modern",
        },
    )
    assert res.status_code == 202
    run_id = res.json()["run_id"]
    poll = client.get(f"/api/ats/runs/{run_id}").json()
    assert poll["status"] == "done"
    assert poll["option_type"] == "jd_paste"
    assert poll["resume"]["contact"]["name"] == "Jane Doe"
    with Session() as s:
        row = s.execute(select(TailorRun).where(TailorRun.run_id == run_id)).scalar_one()
        assert row.option_type == "jd_paste" and row.format_selection == "modern"


def test_upload_docx_then_inject_and_download(api):
    client, Session = api
    docx = _docx_with_runs(["Built Python data pipelines at Acme."])
    up = client.post(
        "/api/ats/parse-upload",
        files={
            "file": (
                "resume.docx",
                docx,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert up.status_code == 200 and up.json()["kind"] == "docx"
    upload_id = up.json()["upload_id"]

    gen = client.post(
        "/api/ats/generate",
        json={
            "option_type": "upload_docx",
            "upload_id": upload_id,
            "jd_text": "Kafka role. " * 20,
            "questions": {},
            "format": "plain",
        },
    )
    assert gen.status_code == 202
    poll = client.get(f"/api/ats/runs/{upload_id}").json()
    assert poll["status"] == "done"
    # Demo mode → no edits computed; diff present with empty applied list.
    assert poll["diff"] is not None
    # Download returns a valid DOCX (original, no edits in demo mode).
    dl = client.post(f"/api/ats/runs/{upload_id}/download-docx", json={})
    assert dl.status_code == 200 and dl.content[:2] == b"PK"


def test_parse_upload_pdf_returns_fallback_kind(api):
    client, _ = api
    res = client.post(
        "/api/ats/parse-upload", files={"file": ("r.pdf", b"%PDF-1.4 fake", "application/pdf")}
    )
    assert res.status_code == 200 and res.json()["kind"] == "pdf"


def test_parse_upload_rejects_other_types(api):
    client, _ = api
    res = client.post("/api/ats/parse-upload", files={"file": ("r.txt", b"hello", "text/plain")})
    assert res.status_code == 415


def test_run_status_404_for_unknown(api):
    client, _ = api
    assert client.get("/api/ats/runs/nope").status_code == 404
