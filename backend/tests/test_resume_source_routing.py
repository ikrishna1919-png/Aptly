"""Resume source-of-truth routing.

The saved DOCX (candidate.active_resume_blob) + default_resume_format.source
decide how BOTH entry points tailor:

  source "resume" (match my resume format) → EXISTING in-place docx_inject path
    against the saved DOCX (never the from-scratch generate path).
  source "ai" (let AI choose)             → the from-scratch generate path.

Covers:
  (a) saving a DOCX persists & survives reload; PDF rejected.
  (b) PUT/GET default_resume_format (+source) round-trips.
  (c) source "resume" → /ats/generate routes to docx_inject vs the saved blob;
      apply_docx_edits leaves non-edited paragraphs byte-for-byte unchanged.
  (d) Jobs tailor flow with source "resume" → docx_inject vs the saved blob.
  (e) source "ai" → /ats/generate uses the generate path (not docx_inject).
"""

from __future__ import annotations

import io

import pytest
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, select
from sqlalchemy.orm import sessionmaker

from app.config import Settings
from app.database import Base
from app.models.candidate import Candidate
from app.models.job import Job
from app.models.tailor_run import TailorRun
from app.models.user import User
from app.services import ats, ats_runs

_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _docx_bytes(paragraphs: list[str]) -> bytes:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def api(tmp_path, monkeypatch):
    from app import config as config_module
    from app.api.auth import get_current_user
    from app.config import get_settings
    from app.database import get_db
    from app.main import app

    engine = create_engine(
        f"sqlite+pysqlite:///{tmp_path / 'rt.db'}",
        future=True,
        connect_args={"check_same_thread": False},
    )
    Base.metadata.create_all(engine)
    Session = sessionmaker(bind=engine, future=True)
    with Session() as s:
        s.add(User(id=1, google_subject_id="g", email="t@e.com", name="T"))
        s.commit()
    user = type("U", (), {"id": 1, "email": "t@e.com", "name": "T"})()

    def override_db():
        with Session() as s:
            yield s

    # The background helpers open their OWN SessionLocal — point them at the
    # test DB, and DON'T launch the worker (we assert on the run row it parks,
    # not on a real keyword-inject/LLM call).
    monkeypatch.setattr(ats_runs, "SessionLocal", Session)
    monkeypatch.setattr(ats_runs, "_launch", lambda target, args: None)

    app.dependency_overrides[get_db] = override_db
    app.dependency_overrides[get_settings] = lambda: Settings(
        DATABASE_URL="x", ADMIN_TOKEN="t", ANTHROPIC_API_KEY=""
    )
    app.dependency_overrides[get_current_user] = lambda: user
    config_module.get_settings.cache_clear()
    try:
        yield TestClient(app), Session
    finally:
        app.dependency_overrides.clear()
        config_module.get_settings.cache_clear()


def _save_resume(client, paras=("Python developer.", "Built data pipelines.")) -> bytes:
    blob = _docx_bytes(list(paras))
    r = client.post("/api/profile/active-resume", files={"file": ("me.docx", blob, _DOCX)})
    assert r.status_code == 200
    return blob


def _set_source(client, source: str, fmt: str = "modern"):
    r = client.post(
        "/api/ats/default-format", json={"kind": "resume", "format": fmt, "source": source}
    )
    assert r.status_code == 200
    return r.json()


# ── (a) ──────────────────────────────────────────────────────────────────────


def test_saving_docx_persists_survives_reload_pdf_rejected(api):
    client, Session = api
    blob = _save_resume(client)
    # Survives reload: a fresh DB session still sees the blob.
    with Session() as s:
        cand = s.execute(select(Candidate).where(Candidate.user_id == 1)).scalar_one()
        assert cand.active_resume_blob == blob
        assert cand.active_resume_content_type == _DOCX
    # PDF rejected with a clear 415.
    pdf = client.post(
        "/api/profile/active-resume", files={"file": ("me.pdf", b"%PDF-1.4", "application/pdf")}
    )
    assert pdf.status_code == 415
    assert ".docx" in pdf.json()["detail"]


# ── (b) ──────────────────────────────────────────────────────────────────────


def test_default_resume_format_source_round_trips(api):
    client, _ = api
    assert _set_source(client, "resume")["source"] == "resume"
    got = client.get("/api/ats/default-format/resume").json()
    assert got["source"] == "resume" and got["format"] == "modern"
    # Switch to AI.
    assert _set_source(client, "ai", fmt="minimal")["source"] == "ai"
    again = client.get("/api/ats/default-format/resume").json()
    assert again["source"] == "ai" and again["format"] == "minimal"


# ── (c) ──────────────────────────────────────────────────────────────────────


def test_generate_source_resume_routes_to_docx_inject(api):
    client, Session = api
    blob = _save_resume(client)
    _set_source(client, "resume")
    r = client.post(
        "/api/ats/generate", json={"option_type": "jd_paste", "jd_text": "need Kafka and Python"}
    )
    assert r.status_code == 202
    run_id = r.json()["run_id"]
    with Session() as s:
        run = s.execute(select(TailorRun).where(TailorRun.run_id == run_id)).scalar_one()
    # Routed to in-place docx_inject against the SAVED resume — not generate.
    assert run.option_type == "upload_docx"
    assert run.uploaded_docx_blob == blob


def test_generate_source_resume_without_saved_docx_400s(api):
    client, _ = api
    _set_source(client, "resume")  # no resume saved
    r = client.post("/api/ats/generate", json={"option_type": "jd_paste", "jd_text": "x"})
    assert r.status_code == 400


def test_apply_docx_edits_leaves_other_paragraphs_byte_for_byte(api):
    """Option-B guarantee: only the matched run changes; every other paragraph
    serialises byte-for-byte identically."""
    from docx import Document

    blob = _docx_bytes(["Python developer.", "Built data pipelines."])
    edited, applied, _skipped = ats.apply_docx_edits(
        blob, [{"original_text": "Python", "replacement_text": "Python and Kafka"}]
    )
    assert len(applied) == 1
    d0 = Document(io.BytesIO(blob))
    d1 = Document(io.BytesIO(edited))
    assert d1.paragraphs[0].text == "Python and Kafka developer."  # edited
    # The untouched paragraph's XML is identical.
    assert d0.paragraphs[1]._p.xml == d1.paragraphs[1]._p.xml


# ── (d) ──────────────────────────────────────────────────────────────────────


def test_jobs_flow_source_resume_routes_to_docx_inject(api):
    client, Session = api
    blob = _save_resume(client)
    _set_source(client, "resume")
    with Session() as s:
        s.add(
            Job(
                id=1,
                source="greenhouse",
                external_id="x1",
                company="Acme",
                title="Engineer",
                url="https://example.com/x1",
                description="We need Kafka and Python.",
            )
        )
        s.commit()
    r = client.post("/api/tailor/start", json={"job_id": 1, "force": True})
    assert r.status_code == 202
    run_id = r.json()["run_id"]
    with Session() as s:
        run = s.execute(select(TailorRun).where(TailorRun.run_id == run_id)).scalar_one()
    assert run.option_type == "upload_docx"
    assert run.uploaded_docx_blob == blob
    # The Jobs run status reports docx mode (no editable resume JSON).
    out = client.get(f"/api/tailor/runs/{run_id}").json()
    assert out["mode"] == "docx_inject"
    assert out["resume"] is None


# ── (e) ──────────────────────────────────────────────────────────────────────


def test_generate_source_ai_uses_generate_path(api):
    client, Session = api
    _save_resume(client)  # a resume IS saved, but source "ai" must ignore it
    _set_source(client, "ai")
    r = client.post("/api/ats/generate", json={"option_type": "jd_paste", "jd_text": "need Python"})
    assert r.status_code == 202
    run_id = r.json()["run_id"]
    with Session() as s:
        run = s.execute(select(TailorRun).where(TailorRun.run_id == run_id)).scalar_one()
    # From-scratch generate path — NOT docx_inject.
    assert run.option_type == "jd_paste"
